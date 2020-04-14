import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt
doc = docx.Document()

# Get guests list 
guestsTxt = open('guests.txt', 'r')
guests = guestsTxt.readlines()
 
for guestName in guests:
    # Add new paragraph
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # justify center

    # 1st line
    run = para.add_run('It would be a pleasure to have the company of\n')
    run.font.name = 'Brush Script Std'
    run.font.size = Pt(20)

    # 2nd line
    run = para.add_run(guestName)
    run.bold = True
    run.font.name = 'Time News Roman'
    run.font.size = Pt(20)

    # 3rd line
    run = para.add_run('at 11010 Memory Lane on the Evening of\n')
    run.font.name = 'Brush Script Std'
    run.font.size = Pt(20)
    
    # 4th line
    run = para.add_run('April 1st\n')
    run.font.name = 'Time News Roman'
    run.font.size = Pt(20)

    # 5th line
    run = para.add_run('at 7 o\'clock')
    run.font.name = 'Brush Script Std'
    run.font.size = Pt(20)

    # Page break
    run.add_break(WD_BREAK.PAGE)

doc.save('invitation.docx')